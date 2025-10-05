"""
Document Processor
Handles extraction and processing of different file types:
- PDF: Text extraction
- DOCX: Text extraction  
- Images: OCR using Tesseract
- Audio: Transcription using Whisper
"""

import os
from pathlib import Path
from typing import Dict, Any, List
import mimetypes

# PDF processing
from pypdf import PdfReader

# DOCX processing
from docx import Document

# Image processing
from PIL import Image
import pytesseract

# Audio processing
import whisper
import torch

# Text chunking
from llama_index.core import Document as LlamaDocument
from llama_index.core.node_parser import SentenceSplitter


class DocumentProcessor:
    def __init__(self):
        """Initialize document processor with models"""
        self.whisper_model = None  # Lazy load
        self.sentence_splitter = SentenceSplitter(
            chunk_size=512,
            chunk_overlap=50
        )
        
    def _load_whisper(self):
        """Lazy load Whisper model"""
        if self.whisper_model is None:
            print("Loading Whisper model...")
            self.whisper_model = whisper.load_model("base")
        return self.whisper_model
    
    async def process_file(self, file_path: str, original_filename: str, file_id: str) -> Dict[str, Any]:
        """
        Process a file based on its type
        Returns extracted text and metadata
        """
        file_path = Path(file_path)
        mime_type, _ = mimetypes.guess_type(str(file_path))
        file_extension = file_path.suffix.lower()
        
        # Determine file type - prioritize extension over mime type for reliability
        if file_extension == '.pdf':
            return await self._process_pdf(file_path, file_id, original_filename)
        elif file_extension in ['.docx', '.doc']:
            return await self._process_docx(file_path, file_id, original_filename)
        elif file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp', '.tiff']:
            return await self._process_image(file_path, file_id, original_filename)
        elif file_extension in ['.mp3', '.wav', '.ogg', '.m4a', '.flac', '.aac', '.wma']:
            return await self._process_audio(file_path, file_id, original_filename)
        elif mime_type and mime_type.startswith('image/'):
            return await self._process_image(file_path, file_id, original_filename)
        elif mime_type and mime_type.startswith('audio/'):
            return await self._process_audio(file_path, file_id, original_filename)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
    
    async def _process_pdf(self, file_path: Path, file_id: str, original_filename: str) -> Dict[str, Any]:
        """Extract text from PDF with page number citations"""
        try:
            reader = PdfReader(str(file_path))
            text = ""
            page_mappings = []  # Track which character positions belong to which page
            
            for page_num, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text()
                start_pos = len(text)
                text += page_text + "\n\n"
                end_pos = len(text)
                page_mappings.append({
                    "page": page_num,
                    "start": start_pos,
                    "end": end_pos
                })
            
            # Chunk the text with page number tracking
            chunks = self._chunk_text_with_citations(text, page_mappings, "pdf")
            
            return {
                "id": file_id,
                "filename": original_filename,
                "type": "pdf",
                "text": text,
                "chunks": len(chunks),
                "chunk_data": chunks,
                "pages": len(reader.pages),
                "page_mappings": page_mappings
            }
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
    
    async def _process_docx(self, file_path: Path, file_id: str, original_filename: str) -> Dict[str, Any]:
        """Extract text from DOCX"""
        try:
            doc = Document(str(file_path))
            text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Chunk the text
            chunks = self._chunk_text(text)
            
            return {
                "id": file_id,
                "filename": original_filename,
                "type": "docx",
                "text": text,
                "chunks": len(chunks),
                "chunk_data": chunks,
                "paragraphs": len(doc.paragraphs)
            }
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")
    
    async def _process_image(self, file_path: Path, file_id: str, original_filename: str) -> Dict[str, Any]:
        """Extract text from image using OCR"""
        try:
            image = Image.open(file_path)
            
            # Perform OCR
            text = pytesseract.image_to_string(image)
            
            # Chunk the text
            chunks = self._chunk_text(text) if text.strip() else []
            
            return {
                "id": file_id,
                "filename": original_filename,
                "type": "image",
                "text": text,
                "chunks": len(chunks),
                "chunk_data": chunks,
                "dimensions": image.size
            }
        except Exception as e:
            raise Exception(f"Error processing image: {str(e)}")
    
    async def _process_audio(self, file_path: Path, file_id: str, original_filename: str) -> Dict[str, Any]:
        """Transcribe audio using Whisper with timestamp tracking"""
        try:
            model = self._load_whisper()
            
            # Whisper natively supports most audio formats through ffmpeg
            # including WAV, OGG, MP3, M4A, FLAC, AAC, etc.
            print(f"Transcribing audio file: {original_filename} ({file_path.suffix})")
            
            # Transcribe with word-level timestamps
            result = model.transcribe(str(file_path), word_timestamps=True)
            text = result["text"]
            segments = result.get("segments", [])
            
            # Create timestamp mappings from segments
            timestamp_mappings = []
            for segment in segments:
                timestamp_mappings.append({
                    "start_time": segment.get("start", 0),
                    "end_time": segment.get("end", 0),
                    "text": segment.get("text", ""),
                    "start_char": segment.get("start_char", 0),
                    "end_char": segment.get("end_char", 0)
                })
            
            # Chunk the text with timestamp tracking
            chunks = self._chunk_text_with_timestamps(text, timestamp_mappings, segments)
            
            return {
                "id": file_id,
                "filename": original_filename,
                "type": "audio",
                "text": text,
                "chunks": len(chunks),
                "chunk_data": chunks,
                "language": result.get("language"),
                "segments": len(segments),
                "timestamp_mappings": timestamp_mappings,
                "duration": segments[-1].get("end", 0) if segments else 0
            }
        except Exception as e:
            raise Exception(f"Error processing audio: {str(e)}")
    
    def _chunk_text(self, text: str) -> List[Dict[str, Any]]:
        """
        Split text into chunks for embedding
        """
        if not text or not text.strip():
            return []
        
        # Create LlamaIndex document
        document = LlamaDocument(text=text)
        
        # Split into nodes
        nodes = self.sentence_splitter.get_nodes_from_documents([document])
        
        # Convert to dict format
        chunks = []
        for i, node in enumerate(nodes):
            chunks.append({
                "chunk_id": i,
                "text": node.get_content(),
                "start_idx": node.start_char_idx,
                "end_idx": node.end_char_idx
            })
        
        return chunks
    
    def _chunk_text_with_citations(
        self, 
        text: str, 
        page_mappings: List[Dict[str, Any]], 
        doc_type: str
    ) -> List[Dict[str, Any]]:
        """
        Split text into chunks and track which page/section each chunk comes from
        """
        if not text or not text.strip():
            return []
        
        # Create LlamaIndex document
        document = LlamaDocument(text=text)
        
        # Split into nodes
        nodes = self.sentence_splitter.get_nodes_from_documents([document])
        
        # Convert to dict format with citation info
        chunks = []
        for i, node in enumerate(nodes):
            chunk_start = node.start_char_idx
            chunk_end = node.end_char_idx
            
            # Find which page(s) this chunk spans
            pages = []
            for mapping in page_mappings:
                # Check if chunk overlaps with this page
                if not (chunk_end < mapping["start"] or chunk_start > mapping["end"]):
                    pages.append(mapping["page"])
            
            chunks.append({
                "chunk_id": i,
                "text": node.get_content(),
                "start_idx": chunk_start,
                "end_idx": chunk_end,
                "citation": {
                    "type": doc_type,
                    "pages": pages if pages else [1],  # Default to page 1 if no match
                    "page_range": f"p. {pages[0]}" if len(pages) == 1 else f"pp. {pages[0]}-{pages[-1]}" if pages else "p. 1"
                }
            })
        
        return chunks
    
    def _chunk_text_with_timestamps(
        self, 
        text: str, 
        timestamp_mappings: List[Dict[str, Any]],
        segments: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """
        Split audio transcription into chunks and track timestamps
        """
        if not text or not text.strip():
            return []
        
        # Create LlamaIndex document
        document = LlamaDocument(text=text)
        
        # Split into nodes
        nodes = self.sentence_splitter.get_nodes_from_documents([document])
        
        # Convert to dict format with timestamp info
        chunks = []
        for i, node in enumerate(nodes):
            chunk_text = node.get_content()
            chunk_start = node.start_char_idx
            chunk_end = node.end_char_idx
            
            # Find which segments this chunk overlaps with
            relevant_segments = []
            for segment in segments:
                seg_text = segment.get("text", "")
                # Check if segment text appears in chunk
                if seg_text.strip() in chunk_text or chunk_text in seg_text:
                    relevant_segments.append(segment)
            
            # Calculate time range
            if relevant_segments:
                start_time = relevant_segments[0].get("start", 0)
                end_time = relevant_segments[-1].get("end", 0)
            else:
                # Estimate based on position in text
                text_position_ratio = chunk_start / len(text) if len(text) > 0 else 0
                total_duration = segments[-1].get("end", 0) if segments else 0
                start_time = text_position_ratio * total_duration
                end_time = start_time + 30  # Estimate 30 seconds per chunk
            
            chunks.append({
                "chunk_id": i,
                "text": chunk_text,
                "start_idx": chunk_start,
                "end_idx": chunk_end,
                "citation": {
                    "type": "audio",
                    "start_time": start_time,
                    "end_time": end_time,
                    "timestamp": self._format_timestamp(start_time),
                    "timestamp_range": f"{self._format_timestamp(start_time)} - {self._format_timestamp(end_time)}"
                }
            })
        
        return chunks
    
    def _format_timestamp(self, seconds: float) -> str:
        """Format seconds into MM:SS or HH:MM:SS"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        
        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"
        else:
            return f"{minutes:02d}:{secs:02d}"
